from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(r"D:\IT\dental-clinic-management")
SRC = ROOT / "HuynhHuuDuy-DC24V7X307-TieuLuan-working.docx"
OUT = ROOT / "HuynhHuuDuy-DC24V7X307-TieuLuan-HoanThien.docx"
IMG = ROOT / "frontend" / "public" / "images"


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_run_font(run, bold=False, italic=False, size=13):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_title(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, bold=True, size=15 if level == 1 else 13)
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"- {item}")
        set_run_font(r)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        run = hdr[idx].paragraphs[0].add_run(text)
        set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            run = cells[idx].paragraphs[0].add_run(str(text))
            set_run_font(run, size=12)
    doc.add_paragraph()
    return table


def add_image(doc, file_name, caption, width=5.8):
    path = IMG / file_name
    if not path.exists():
        add_para(doc, f"[Cần bổ sung hình: {file_name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, italic=True, size=12)


def clean_tail_template(doc):
    # Xoa phan huong dan mau con du o cuoi file.
    start = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == "1. Trình bày luận văn":
            start = i
            break
    if start is not None:
        for p in list(doc.paragraphs[start:]):
            remove_paragraph(p)

    # Xoa cac dong placeholder "Bo sung hinh..." neu con nam trong chuong 3.
    for p in list(doc.paragraphs):
        if p.text.strip().lower().startswith("bổ sung hình"):
            remove_paragraph(p)


def build_report():
    doc = Document(SRC)
    clean_tail_template(doc)

    doc.add_page_break()
    add_title(doc, "CHƯƠNG 4. ĐÁNH GIÁ KIỂM THỬ VÀ GIỚI THIỆU CHƯƠNG TRÌNH")

    add_title(doc, "4.1 Mục tiêu kiểm thử", level=2)
    add_para(
        doc,
        "Mục tiêu của chương này là đánh giá mức độ đáp ứng yêu cầu của hệ thống quản lý phòng khám nha khoa V sau khi cài đặt. "
        "Quá trình kiểm thử tập trung vào các luồng chính gồm: đặt lịch khám, quản lý lịch hẹn, quản lý bệnh nhân, hồ sơ điều trị, hóa đơn, đánh giá dịch vụ, tìm kiếm, chatbot tư vấn và khả năng hiển thị trên thiết bị di động."
    )
    add_para(
        doc,
        "Ngoài kiểm thử chức năng, đề tài cũng kiểm tra trải nghiệm sử dụng ở các vai trò khác nhau: khách hàng, quản trị viên/lễ tân và nha sĩ. "
        "Các trường hợp kiểm thử được xây dựng theo tình huống thực tế của một phòng khám nha khoa, đặc biệt là tình huống trùng lịch, bác sĩ bận và khách vãng lai đặt lịch không cần tài khoản."
    )

    add_title(doc, "4.2 Môi trường kiểm thử", level=2)
    add_table(
        doc,
        ["Thành phần", "Công cụ / môi trường sử dụng"],
        [
            ["Frontend", "ReactJS, Vite, trình duyệt Chrome và trình duyệt trên điện thoại"],
            ["Backend", "NodeJS, ExpressJS"],
            ["Cơ sở dữ liệu", "PostgreSQL khi chạy cục bộ, Neon PostgreSQL khi triển khai web thật"],
            ["Triển khai", "Frontend triển khai trên Vercel, backend triển khai trên Render"],
            ["Kiểm thử thủ công", "Đăng nhập theo từng vai trò, đặt lịch, tạo dữ liệu, kiểm tra kết quả trên giao diện"],
        ],
    )

    add_title(doc, "4.3 Kiểm thử chức năng chính", level=2)
    add_table(
        doc,
        ["STT", "Chức năng", "Trường hợp kiểm thử", "Kết quả mong đợi", "Đánh giá"],
        [
            [1, "Đăng ký", "Khách hàng nhập tên đăng nhập, mật khẩu, số điện thoại", "Tạo tài khoản khách hàng và có thể đăng nhập", "Đạt"],
            [2, "Đăng nhập", "Đăng nhập bằng tài khoản admin, nha sĩ, khách hàng", "Hệ thống phân quyền đúng giao diện theo vai trò", "Đạt"],
            [3, "Quên mật khẩu", "Nhập thông tin tài khoản và đặt lại mật khẩu", "Người dùng đổi được mật khẩu để tiếp tục sử dụng", "Đạt"],
            [4, "Đặt lịch", "Khách vãng lai đặt lịch không cần tài khoản", "Tạo lịch hẹn chờ xác nhận và lưu thông tin khách", "Đạt"],
            [5, "Quản lý lịch hẹn", "Admin xác nhận, hủy hoặc hoàn tất lịch hẹn", "Trạng thái lịch thay đổi và hiển thị đúng cho người dùng", "Đạt"],
            [6, "Lịch bận nha sĩ", "Nha sĩ tạo lịch bận/nghỉ", "Admin thấy thông báo và tránh phân công trùng lịch", "Đạt"],
            [7, "Hồ sơ điều trị", "Nha sĩ cập nhật chẩn đoán, nội dung điều trị, ảnh đính kèm và tái khám", "Khách hàng xem được kết quả khám trong tài khoản", "Đạt"],
            [8, "Hóa đơn", "Admin lập hóa đơn theo khách hàng và dịch vụ", "Hệ thống lưu hóa đơn, chi tiết hóa đơn và hỗ trợ in phiếu", "Đạt"],
            [9, "Đánh giá", "Khách hàng đánh giá sau khi hoàn tất lịch khám", "Chỉ lịch đã hoàn tất mới được đánh giá", "Đạt"],
            [10, "Tìm kiếm", "Tìm dịch vụ, bảng giá, bác sĩ, bài thông tin", "Trả về nhóm kết quả phù hợp với từ khóa", "Đạt"],
            [11, "Chatbot AI", "Khách hỏi kiến thức nha khoa, dịch vụ, chi phí tham khảo", "Chatbot trả lời định hướng, có cảnh báo không thay thế bác sĩ", "Đạt"],
        ],
    )

    add_title(doc, "4.4 Kiểm thử luồng đặt lịch khám", level=2)
    add_para(
        doc,
        "Đặt lịch là chức năng quan trọng nhất của hệ thống vì liên quan trực tiếp đến khách hàng, lễ tân và nha sĩ. "
        "Hệ thống cho phép khách đặt lịch có hoặc không có tài khoản. Nếu khách không chọn nha sĩ cụ thể, lịch được đưa vào trạng thái chờ phân công để admin/lễ tân chọn nha sĩ phù hợp."
    )
    add_table(
        doc,
        ["Tình huống", "Cách xử lý của hệ thống"],
        [
            ["Khách chọn nha sĩ cụ thể và khung giờ còn trống", "Hệ thống tạo lịch hẹn cho nha sĩ đó và chờ admin xác nhận."],
            ["Khách chọn nha sĩ cụ thể nhưng nha sĩ đã có lịch cùng giờ", "Hệ thống không cho đặt trùng và yêu cầu chọn khung giờ khác."],
            ["Khách không chọn nha sĩ cụ thể", "Hệ thống tạo lịch chờ phân công để admin xem nha sĩ nào còn trống rồi gán sau."],
            ["Nhiều khách cùng đặt một khung giờ nhưng không chọn nha sĩ", "Admin có thể phân bổ cho nhiều nha sĩ khác nhau nếu phòng khám còn đủ bác sĩ và ghế điều trị."],
            ["Tất cả nha sĩ đều bận trong khung giờ", "Hệ thống hiển thị các giờ còn trống khác để khách chọn lại hoặc admin liên hệ tư vấn đổi lịch."],
            ["Nha sĩ tạo lịch bận/nghỉ", "Admin nhận thông báo và tránh phân công lịch vào khoảng thời gian đó."],
        ],
    )
    add_para(
        doc,
        "Về mặt nghiệp vụ, hệ thống không tự động kết luận toàn bộ ngày đã hết lịch khi chỉ một khung giờ bị trùng. "
        "Thay vào đó, giao diện hiển thị các khung giờ còn khả dụng để khách lựa chọn, giúp giảm tình trạng khách phải thử từng giờ thủ công."
    )

    add_title(doc, "4.5 Kiểm thử chatbot AI và tìm kiếm", level=2)
    add_para(
        doc,
        "Chatbot được xây dựng với vai trò trợ lý tư vấn nha khoa ban đầu. Nội dung trả lời tập trung vào kiến thức phổ biến như đau răng, sâu răng, viêm nướu, viêm tủy, nha chu, răng ê buốt, phục hình răng, implant, niềng răng, tẩy trắng răng và chăm sóc sau điều trị. "
        "Chatbot không đưa ra chẩn đoán thay bác sĩ, mà chỉ giải thích theo ngôn ngữ dễ hiểu và khuyến nghị khách đặt lịch khi có dấu hiệu bất thường."
    )
    add_para(
        doc,
        "Đối với tìm kiếm, hệ thống gom dữ liệu từ nhiều nhóm như dịch vụ, bảng giá, bác sĩ, ưu đãi và thông tin nha khoa. Kết quả được ưu tiên theo mức độ liên quan để khách hàng không phải đọc quá nhiều thông tin không cần thiết."
    )

    add_title(doc, "4.6 Giới thiệu chương trình", level=2)
    add_para(
        doc,
        "Hệ thống được chia thành ba nhóm giao diện chính: giao diện công khai dành cho khách truy cập, giao diện quản trị dành cho admin/lễ tân và giao diện làm việc dành cho nha sĩ."
    )
    add_image(doc, "clinic-hero-care.png", "Hình 4.1. Giao diện trang chủ Nha khoa V", width=5.9)
    add_para(
        doc,
        "Trang chủ cung cấp thông tin tổng quan về phòng khám, dịch vụ, đội ngũ bác sĩ, cơ sở vật chất, ưu đãi, lịch làm việc, bảng giá và nút đặt lịch nhanh. "
        "Thiết kế hướng đến khách hàng phổ thông nên ưu tiên nội dung dễ đọc, rõ luồng hành động và hỗ trợ tốt trên thiết bị di động."
    )
    add_image(doc, "service-implant-consultation.png", "Hình 4.2. Minh họa nhóm dịch vụ nha khoa", width=5.4)
    add_para(
        doc,
        "Trang dịch vụ giúp khách hàng tìm hiểu trước về các nhóm điều trị như nha khoa tổng quát, implant, chỉnh nha, nha khoa thẩm mỹ và nha khoa trẻ em. "
        "Mỗi dịch vụ có trang chi tiết riêng để giải thích đối tượng phù hợp, quy trình, lưu ý và hướng đặt lịch."
    )
    add_image(doc, "clinic-story-equipment.png", "Hình 4.3. Trang giới thiệu cơ sở vật chất và thiết bị", width=5.9)
    add_para(
        doc,
        "Phần cơ sở vật chất giới thiệu các thiết bị hỗ trợ điều trị như ghế nha khoa RunTour, máy CBCT 3 in 1 Hyperion X5, máy Scan Shining 3D và nồi hấp Vacuclave MELAG 323. "
        "Thông tin này giúp khách hàng hiểu phòng khám có đầu tư thiết bị hỗ trợ chẩn đoán, lấy dấu kỹ thuật số và kiểm soát vô trùng dụng cụ."
    )
    add_image(doc, "promo-implant-summer.png", "Hình 4.4. Minh họa chương trình ưu đãi", width=5.2)
    add_para(
        doc,
        "Chương trình ưu đãi được tách thành các trang riêng để tránh làm loãng nội dung trang chủ. Khách hàng có thể xem chi tiết ưu đãi implant, chỉnh nha hoặc điều trị tổng quát trước khi đặt lịch tư vấn."
    )
    add_image(doc, "home-doctor-01.png", "Hình 4.5. Minh họa hồ sơ bác sĩ", width=4.6)
    add_para(
        doc,
        "Thông tin bác sĩ được hiển thị theo hồ sơ chuyên môn, gồm chuyên khoa, kinh nghiệm, thế mạnh điều trị và thông tin liên hệ. Khi admin cập nhật bác sĩ trong hệ thống, danh sách bác sĩ có thể được dùng để phục vụ đặt lịch và giới thiệu đội ngũ trên giao diện công khai."
    )

    add_title(doc, "4.7 Kiểm thử triển khai web thật", level=2)
    add_para(
        doc,
        "Ngoài môi trường localhost, hệ thống đã được triển khai thử nghiệm lên môi trường web thật để người dùng có thể truy cập bằng đường dẫn công khai. "
        "Frontend được triển khai trên Vercel, backend được triển khai trên Render và cơ sở dữ liệu sử dụng Neon PostgreSQL. Việc triển khai này giúp kiểm tra khả năng hoạt động của hệ thống khi tách frontend, backend và database thành các dịch vụ riêng."
    )
    add_table(
        doc,
        ["Thành phần", "Nền tảng triển khai", "Vai trò"],
        [
            ["Frontend", "Vercel", "Chạy giao diện React, nhận thao tác người dùng"],
            ["Backend", "Render", "Cung cấp API xử lý đăng nhập, đặt lịch, quản lý dữ liệu"],
            ["Database", "Neon PostgreSQL", "Lưu dữ liệu tài khoản, bệnh nhân, lịch hẹn, hóa đơn, hồ sơ điều trị"],
        ],
    )
    add_para(
        doc,
        "Khi chạy demo tại lớp, hệ thống vẫn có thể chạy ở localhost để dễ kiểm soát dữ liệu và thao tác kỹ thuật. "
        "Đường dẫn web thật được dùng như một minh chứng rằng hệ thống có khả năng triển khai ngoài môi trường máy cá nhân."
    )

    doc.add_page_break()
    add_title(doc, "CHƯƠNG 5. KẾT LUẬN")

    add_title(doc, "5.1 Kết quả đạt được", level=2)
    add_para(
        doc,
        "Sau quá trình phân tích, thiết kế và cài đặt, đề tài đã xây dựng được hệ thống quản lý phòng khám nha khoa V theo hướng web fullstack. "
        "Hệ thống đáp ứng các nghiệp vụ chính của phòng khám, đồng thời bổ sung các chức năng hỗ trợ khách hàng trước và sau khi đặt lịch."
    )
    add_bullets(
        doc,
        [
            "Xây dựng giao diện công khai gồm trang chủ, giới thiệu phòng khám, dịch vụ, bảng giá, đội ngũ bác sĩ, cơ sở vật chất, ưu đãi, đặt lịch và chatbot tư vấn.",
            "Xây dựng chức năng đặt lịch cho khách có tài khoản và khách vãng lai, hỗ trợ kiểm tra trùng lịch và theo dõi trạng thái lịch hẹn.",
            "Xây dựng giao diện quản trị để quản lý khách hàng, nha sĩ, dịch vụ, lịch hẹn, hóa đơn, đánh giá và thông báo lịch bận.",
            "Xây dựng giao diện nha sĩ để theo dõi lịch khám, tạo lịch bận và cập nhật hồ sơ điều trị cho bệnh nhân.",
            "Tích hợp chatbot tư vấn nha khoa cơ bản, giúp khách hàng hiểu thông tin ban đầu về bệnh lý răng miệng, dịch vụ và chi phí tham khảo.",
            "Triển khai thử nghiệm hệ thống lên môi trường web thật bằng Vercel, Render và Neon PostgreSQL.",
        ],
    )

    add_title(doc, "5.2 Hạn chế", level=2)
    add_para(
        doc,
        "Do thời gian thực hiện đề tài có giới hạn, hệ thống vẫn còn một số hạn chế nhất định. Chatbot chỉ dừng ở mức tư vấn thông tin tham khảo, chưa thể thay thế bác sĩ trong việc chẩn đoán hoặc đưa ra phác đồ điều trị. "
        "Một số dữ liệu demo như bác sĩ, dịch vụ và chương trình ưu đãi được xây dựng phục vụ mục tiêu trình bày tiểu luận, chưa phản ánh đầy đủ toàn bộ hoạt động của một phòng khám nha khoa quy mô lớn."
    )
    add_para(
        doc,
        "Bên cạnh đó, phiên bản triển khai miễn phí trên các nền tảng cloud có thể bị giới hạn tài nguyên. Backend miễn phí có thể khởi động chậm sau một thời gian không truy cập, vì vậy tốc độ phản hồi trên web thật có thể chưa ổn định như khi chạy ở localhost."
    )

    add_title(doc, "5.3 Hướng phát triển", level=2)
    add_bullets(
        doc,
        [
            "Nâng cấp chatbot AI theo hướng kết hợp tri thức nội bộ và mô hình AI bên ngoài để trả lời tự nhiên, đúng trọng tâm và bao phủ nhiều kiến thức nha khoa hơn.",
            "Bổ sung chức năng tra cứu bảo hành implant, răng sứ hoặc phục hình theo số điện thoại/mã bảo hành cho khách không cần đăng nhập.",
            "Tích hợp gửi thông báo qua Zalo, SMS hoặc email để nhắc lịch khám, tái khám và thông báo thay đổi lịch.",
            "Tăng cường bảo mật bằng cách mã hóa mật khẩu mạnh hơn, giới hạn số lần đăng nhập sai, kiểm tra dữ liệu đầu vào và phân quyền API chặt chẽ hơn.",
            "Tối ưu hiệu năng, kiểm thử tải và nâng cấp hạ tầng khi số lượng người truy cập hoặc số lượng phòng khám tăng lên.",
            "Mở rộng hệ thống theo mô hình nhiều chi nhánh, nhiều ghế điều trị và nhiều nhóm nhân sự hơn.",
        ],
    )

    add_title(doc, "5.4 Kết luận chung", level=2)
    add_para(
        doc,
        "Đề tài “Xây dựng hệ thống quản lý phòng khám nha khoa thông minh” đã hoàn thành mục tiêu xây dựng một hệ thống web hỗ trợ quản lý phòng khám nha khoa và nâng cao trải nghiệm khách hàng. "
        "Hệ thống không chỉ hỗ trợ các nghiệp vụ cơ bản như đăng nhập, đặt lịch, quản lý bệnh nhân, nha sĩ, dịch vụ, hóa đơn và hồ sơ điều trị, mà còn bổ sung các thành phần phù hợp với nhu cầu hiện nay như chatbot tư vấn, tìm kiếm thông tin, giao diện giới thiệu dịch vụ và triển khai web thật."
    )
    add_para(
        doc,
        "Thông qua quá trình thực hiện đề tài, người thực hiện có cơ hội vận dụng kiến thức về phân tích yêu cầu, thiết kế cơ sở dữ liệu, xây dựng API, phát triển giao diện ReactJS, kiểm thử chức năng và triển khai hệ thống. "
        "Kết quả đạt được là nền tảng để tiếp tục phát triển thành một hệ thống quản lý phòng khám hoàn thiện hơn trong thực tế."
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_report())
