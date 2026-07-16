from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\IT\dental-clinic-management")
DOCX = ROOT / "HuynhHuuDuy-DC24V7X307-TieuLuan-HoanThien.docx"
OUT = ROOT / "HuynhHuuDuy-DC24V7X307-TieuLuan-HoanThien.docx"
ASSET_DIR = ROOT / "report_assets"
ASSET_DIR.mkdir(exist_ok=True)


def font(size=26, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return ImageFont.truetype(c, size)
    return ImageFont.load_default()


def draw_box(draw, xy, text, fill="#fffaf2", outline="#c98b2d", title=False):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    lines = wrap_text(text, 24 if title else 26)
    f = font(24 if title else 22, bold=title)
    total_h = len(lines) * 30
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#18243a", font=f)
        y += 30


def wrap_text(text, limit):
    words = text.split()
    lines, cur = [], ""
    for w in words:
        if len(cur + " " + w) <= limit:
            cur = (cur + " " + w).strip()
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def arrow(draw, start, end):
    draw.line([start, end], fill="#a26013", width=4)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)]
    else:
        pts = [(x2, y2), (x2 + 18, y2 - 10), (x2 + 18, y2 + 10)]
    draw.polygon(pts, fill="#a26013")


def save_architecture():
    img = Image.new("RGB", (1200, 620), "#fff8ef")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Kien truc tong the he thong", fill="#8a4b0f", font=font(32, True))
    boxes = [
        ((60, 170, 300, 310), "Nguoi dung\nKhach hang / Admin / Nha si"),
        ((390, 150, 650, 330), "Frontend ReactJS\nGiao dien + Router + Axios"),
        ((740, 150, 980, 330), "Backend NodeJS\nREST API + JWT + Nghiep vu"),
        ((390, 420, 650, 560), "PostgreSQL\nUsers, patients, appointments..."),
        ((740, 420, 980, 560), "Gemini API\nHo tro chatbot AI"),
    ]
    for xy, text in boxes:
        draw_box(d, xy, text)
    arrow(d, (300, 240), (390, 240))
    arrow(d, (650, 240), (740, 240))
    arrow(d, (860, 330), (860, 420))
    arrow(d, (740, 490), (650, 490))
    img.save(ASSET_DIR / "chapter3-architecture.png")


def save_usecase():
    img = Image.new("RGB", (1200, 700), "#fff8ef")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Use case tong quat theo vai tro", fill="#8a4b0f", font=font(32, True))
    cols = [
        ("Khach hang", ["Dang ky / dang nhap", "Dat lich kham", "Xem lich va ket qua", "Danh gia dich vu", "Chatbot tu van"]),
        ("Admin / Le tan", ["Quan ly lich hen", "Quan ly benh nhan", "Quan ly nha si", "Lap hoa don", "Theo doi thong ke"]),
        ("Nha si", ["Xem lich kham", "Tao lich ban", "Cap nhat ho so", "De xuat tai kham", "Xem benh nhan"]),
    ]
    x = 70
    for title, items in cols:
        draw_box(d, (x, 130, x + 310, 210), title, fill="#f6d59d", title=True)
        y = 245
        for item in items:
            draw_box(d, (x, y, x + 310, y + 65), item)
            y += 82
        x += 380
    img.save(ASSET_DIR / "chapter3-usecase.png")


def save_erd():
    img = Image.new("RGB", (1300, 760), "#fff8ef")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Mo hinh du lieu chinh", fill="#8a4b0f", font=font(32, True))
    boxes = {
        "users": (70, 130, 270, 230),
        "patients": (360, 130, 560, 230),
        "dentists": (650, 130, 850, 230),
        "appointments": (360, 330, 600, 445),
        "services": (70, 330, 270, 430),
        "medical_records": (700, 330, 980, 445),
        "invoices": (360, 560, 570, 660),
        "invoice_details": (670, 560, 930, 660),
        "reviews": (1010, 330, 1220, 430),
        "chatbot_logs": (1010, 560, 1220, 660),
    }
    for name, xy in boxes.items():
        draw_box(d, xy, name, title=True)
    for a, b in [
        ("users", "patients"),
        ("users", "dentists"),
        ("patients", "appointments"),
        ("dentists", "appointments"),
        ("services", "appointments"),
        ("appointments", "medical_records"),
        ("patients", "invoices"),
        ("invoices", "invoice_details"),
        ("appointments", "reviews"),
    ]:
        x1 = (boxes[a][0] + boxes[a][2]) // 2
        y1 = (boxes[a][1] + boxes[a][3]) // 2
        x2 = (boxes[b][0] + boxes[b][2]) // 2
        y2 = (boxes[b][1] + boxes[b][3]) // 2
        d.line([(x1, y1), (x2, y2)], fill="#a26013", width=3)
    img.save(ASSET_DIR / "chapter3-erd.png")


def save_booking_flow():
    img = Image.new("RGB", (1300, 520), "#fff8ef")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Luong xu ly dat lich", fill="#8a4b0f", font=font(32, True))
    steps = [
        "Khach nhap thong tin",
        "Chon dich vu, ngay gio",
        "Kiem tra khung gio",
        "Tao lich cho xac nhan",
        "Admin phan cong / xac nhan",
        "Nha si nhan lich kham",
    ]
    x = 45
    for step in steps:
        draw_box(d, (x, 185, x + 185, 310), step)
        if step != steps[-1]:
            arrow(d, (x + 185, 248), (x + 230, 248))
        x += 210
    img.save(ASSET_DIR / "chapter3-booking-flow.png")


def save_chatbot_flow():
    img = Image.new("RGB", (1300, 560), "#fff8ef")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Luong xu ly chatbot AI", fill="#8a4b0f", font=font(32, True))
    steps = [
        "Nguoi dung hoi",
        "Chuan hoa cau hoi\nkhong dau / viet tat",
        "Tim tri thuc noi bo",
        "Goi Gemini API\nkhi can bo sung",
        "Tra loi + canh bao\nkhong thay bac si",
        "Luu chatbot_logs",
    ]
    x = 45
    for step in steps:
        draw_box(d, (x, 190, x + 185, 335), step)
        if step != steps[-1]:
            arrow(d, (x + 185, 262), (x + 230, 262))
        x += 210
    img.save(ASSET_DIR / "chapter3-chatbot-flow.png")


def set_run_font(run, bold=False, italic=False, size=13):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_before(ref, text="", bold=False, italic=False, size=13, align=None):
    p = ref.insert_paragraph_before()
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, bold=bold, italic=italic, size=size)
    return p


def add_picture_before(ref, path, caption):
    p = ref.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(5.8))
    cap = ref.insert_paragraph_before()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, italic=True, size=12)


def remove_block(doc, start_text, end_text):
    paragraphs = list(doc.paragraphs)
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith(start_text))
    end = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith(end_text))
    ref = paragraphs[end]
    for p in paragraphs[start:end]:
        el = p._element
        el.getparent().remove(el)
    return ref


def rewrite_chapter3():
    for fn in [save_architecture, save_usecase, save_erd, save_booking_flow, save_chatbot_flow]:
        fn()

    doc = Document(DOCX)
    ref = remove_block(doc, "CHƯƠNG 3", "CHƯƠNG 4")

    add_before(ref, "CHƯƠNG 3. THIẾT KẾ VÀ CÀI ĐẶT GIẢI PHÁP", bold=True, size=15)
    add_before(ref, "3.1 Định hướng thiết kế hệ thống", bold=True)
    add_before(ref, "Hệ thống được thiết kế theo mô hình web fullstack, tách rõ phần giao diện, phần xử lý nghiệp vụ và phần lưu trữ dữ liệu. Cách tổ chức này giúp chương trình dễ bảo trì, dễ mở rộng và phù hợp với mô hình triển khai thực tế của một website phòng khám nha khoa.")
    add_before(ref, "Frontend chịu trách nhiệm hiển thị giao diện và nhận thao tác của người dùng. Backend tiếp nhận yêu cầu thông qua REST API, kiểm tra dữ liệu, xử lý nghiệp vụ và truy vấn cơ sở dữ liệu PostgreSQL. Các chức năng AI như chatbot được đặt ở backend để dễ kiểm soát nội dung trả lời, lưu lịch sử hỏi đáp và bảo vệ khóa API.")
    add_picture_before(ref, ASSET_DIR / "chapter3-architecture.png", "Hình 3.1. Kiến trúc tổng thể của hệ thống")

    add_before(ref, "3.2 Công nghệ sử dụng", bold=True)
    add_before(ref, "Các công nghệ được lựa chọn dựa trên tiêu chí dễ triển khai, phù hợp với phạm vi tiểu luận và có khả năng mở rộng khi phát triển tiếp.")
    add_before(ref, "- ReactJS và Vite: xây dựng giao diện người dùng, chia nhỏ thành component để dễ quản lý.")
    add_before(ref, "- React Router DOM: điều hướng giữa trang chủ, bảng giá, đặt lịch, tài khoản khách hàng, admin và nha sĩ.")
    add_before(ref, "- Axios: gửi request từ frontend đến backend API.")
    add_before(ref, "- NodeJS và ExpressJS: xây dựng server, định nghĩa route và xử lý logic nghiệp vụ.")
    add_before(ref, "- PostgreSQL: lưu dữ liệu người dùng, bệnh nhân, nha sĩ, lịch hẹn, hồ sơ điều trị, hóa đơn và chatbot logs.")
    add_before(ref, "- JWT: xác thực người dùng và phân quyền theo vai trò admin, nha sĩ, khách hàng.")
    add_before(ref, "- Gemini API: hỗ trợ chatbot trả lời các câu hỏi nha khoa ngoài bộ tri thức nội bộ.")

    add_before(ref, "3.3 Thiết kế chức năng theo vai trò", bold=True)
    add_before(ref, "Hệ thống có ba nhóm người dùng chính: khách hàng, quản trị viên/lễ tân và nha sĩ. Mỗi vai trò được phân quyền riêng để tránh truy cập sai chức năng.")
    add_picture_before(ref, ASSET_DIR / "chapter3-usecase.png", "Hình 3.2. Use case tổng quát theo vai trò")
    add_before(ref, "Khách hàng có thể đăng ký, đăng nhập, đặt lịch khám, xem lịch đã đặt, xem kết quả điều trị, đánh giá dịch vụ sau khi hoàn tất lịch khám và sử dụng chatbot tư vấn. Admin/lễ tân quản lý dữ liệu vận hành như khách hàng, nha sĩ, dịch vụ, lịch hẹn, hóa đơn, đánh giá và dashboard thống kê. Nha sĩ xem lịch khám được phân công, tạo lịch bận và cập nhật hồ sơ điều trị.")

    add_before(ref, "3.4 Thiết kế cơ sở dữ liệu", bold=True)
    add_before(ref, "Cơ sở dữ liệu được thiết kế theo mô hình quan hệ. Các bảng chính gồm users, patients, dentists, services, appointments, medical_records, medical_record_attachments, invoices, invoice_details, reviews, dentist_unavailable_times, page_visits và chatbot_logs.")
    add_before(ref, "Bảng appointments đóng vai trò trung tâm trong luồng đặt lịch vì liên kết bệnh nhân, nha sĩ và dịch vụ. Bảng medical_records lưu kết quả khám và đề xuất tái khám. Bảng invoices và invoice_details tách riêng hóa đơn tổng và chi tiết dịch vụ để phù hợp nghiệp vụ thanh toán thực tế.")
    add_picture_before(ref, ASSET_DIR / "chapter3-erd.png", "Hình 3.3. Mô hình dữ liệu chính của hệ thống")

    add_before(ref, "3.5 Thiết kế luồng đặt lịch khám", bold=True)
    add_before(ref, "Luồng đặt lịch được thiết kế để hỗ trợ cả khách có tài khoản và khách vãng lai. Khi khách chọn ngày giờ, hệ thống kiểm tra lịch đã có và lịch bận của nha sĩ. Nếu khách chọn nha sĩ cụ thể, hệ thống không cho phép tạo lịch trùng cùng nha sĩ trong cùng khung giờ. Nếu khách không chọn nha sĩ, lịch được tạo ở trạng thái chờ phân công để admin/lễ tân gán nha sĩ còn trống.")
    add_picture_before(ref, ASSET_DIR / "chapter3-booking-flow.png", "Hình 3.4. Luồng xử lý đặt lịch khám")

    add_before(ref, "3.6 Thiết kế chatbot AI", bold=True)
    add_before(ref, "Chatbot được thiết kế theo hướng kết hợp tri thức nội bộ và API AI bên ngoài. Trước tiên hệ thống chuẩn hóa câu hỏi của người dùng, bao gồm trường hợp viết không dấu, viết tắt hoặc sai chính tả đơn giản. Sau đó backend tìm câu trả lời phù hợp trong bộ tri thức nha khoa nội bộ. Nếu câu hỏi cần bổ sung hoặc không khớp rõ ràng, hệ thống có thể gọi Gemini API để hỗ trợ tạo câu trả lời tự nhiên hơn.")
    add_before(ref, "Mọi câu trả lời đều phải giữ nguyên nguyên tắc an toàn: chỉ cung cấp thông tin tham khảo, không chẩn đoán thay bác sĩ và khuyến nghị đặt lịch khám khi có dấu hiệu đau kéo dài, sưng, chảy máu, sốt, răng lung lay hoặc nghi ngờ nhiễm trùng.")
    add_picture_before(ref, ASSET_DIR / "chapter3-chatbot-flow.png", "Hình 3.5. Luồng xử lý chatbot AI")

    add_before(ref, "3.7 Thiết kế giao diện", bold=True)
    add_before(ref, "Giao diện được chia thành ba khu vực: website công khai, trang quản trị và trang nha sĩ. Website công khai tập trung vào thông tin dịch vụ, bảng giá, đội ngũ bác sĩ, cơ sở vật chất, ưu đãi và đặt lịch. Trang quản trị tập trung vào dữ liệu vận hành. Trang nha sĩ tập trung vào lịch khám, lịch bận và hồ sơ điều trị.")
    add_before(ref, "Các thành phần giao diện dùng lại như navbar, footer, logo, nút cuộn lên, floating support, layout admin và layout nha sĩ được tách thành component riêng để giảm lặp code.")

    add_before(ref, "3.8 Cài đặt backend", bold=True)
    add_before(ref, "Backend được tổ chức theo mô hình gần với MVC. Thư mục routes khai báo đường dẫn API, controllers xử lý request/response, models truy vấn cơ sở dữ liệu, middlewares kiểm tra đăng nhập và uploadMiddleware xử lý file đính kèm. Cách chia này giúp từng phần có trách nhiệm rõ ràng và dễ giải thích khi bảo vệ đề tài.")
    add_before(ref, "Một số API quan trọng gồm: auth API cho đăng ký/đăng nhập/quên mật khẩu, appointment API cho đặt lịch và xử lý lịch hẹn, dentist API cho quản lý nha sĩ, medical record API cho hồ sơ điều trị, invoice API cho hóa đơn, chatbot API cho tư vấn AI và dashboard API cho thống kê.")

    add_before(ref, "3.9 Cài đặt frontend", bold=True)
    add_before(ref, "Frontend sử dụng ReactJS. Các trang chính được đặt trong thư mục pages, component dùng lại đặt trong components, các layout riêng cho admin và nha sĩ đặt trong layouts. File axiosClient cấu hình base URL để frontend gọi backend, giúp khi deploy chỉ cần đổi biến môi trường VITE_API_BASE_URL.")
    add_before(ref, "Trạng thái đăng nhập được lưu bằng token và thông tin người dùng trong localStorage. Khi người dùng truy cập trang cần quyền, RoleRoute kiểm tra vai trò trước khi cho vào trang tương ứng.")

    add_before(ref, "3.10 Cài đặt và triển khai", bold=True)
    add_before(ref, "Khi chạy cục bộ, backend kết nối PostgreSQL bằng biến môi trường trong file .env, frontend chạy bằng Vite. Khi triển khai web thật, frontend được đưa lên Vercel, backend được đưa lên Render và database sử dụng Neon PostgreSQL. Cách triển khai này giúp tách riêng giao diện, API và dữ liệu, gần với mô hình triển khai thực tế.")

    doc.save(OUT)


if __name__ == "__main__":
    rewrite_chapter3()
    print(OUT)
